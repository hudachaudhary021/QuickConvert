"""
converter.py
------------
Conversion helpers for QuickConvert.

Word -> PDF:
    Reads the .docx with python-docx and re-renders the textual content
    (paragraphs, headings, bullet/numbered lists, basic bold/italic runs,
    and tables) into a PDF using reportlab. This is a pure-Python path
    that works on Termux, Render, and any other host without requiring
    LibreOffice or Microsoft Word to be installed.

PDF -> Word:
    Uses the pdf2docx library, which parses the PDF layout directly and
    rebuilds paragraphs, tables, and images into a native .docx file.
    Also pure-Python, no external binaries required.
"""

import os

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from reportlab.lib.pagesizes import LETTER
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, ListFlowable, ListItem
)
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT

from pdf2docx import Converter as PDF2DocxConverter


class ConversionError(Exception):
    """Raised when a file cannot be validated or converted."""


# --------------------------------------------------------------------------
# Word -> PDF
# --------------------------------------------------------------------------

def _build_pdf_styles():
    styles = getSampleStyleSheet()

    styles.add(ParagraphStyle(
        name="QCHeading1", parent=styles["Heading1"],
        fontSize=20, leading=24, spaceAfter=14, textColor=colors.HexColor("#1b2340"),
    ))
    styles.add(ParagraphStyle(
        name="QCHeading2", parent=styles["Heading2"],
        fontSize=16, leading=20, spaceAfter=12, textColor=colors.HexColor("#1b2340"),
    ))
    styles.add(ParagraphStyle(
        name="QCHeading3", parent=styles["Heading3"],
        fontSize=13, leading=17, spaceAfter=10, textColor=colors.HexColor("#1b2340"),
    ))
    styles.add(ParagraphStyle(
        name="QCBody", parent=styles["Normal"],
        fontSize=11, leading=16, spaceAfter=8, alignment=TA_LEFT,
    ))
    return styles


def _runs_to_html(paragraph):
    """Convert python-docx runs into a reportlab-friendly HTML-ish string,
    preserving bold/italic/underline where present."""
    parts = []
    for run in paragraph.runs:
        text = (run.text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if not text:
            continue
        if run.bold:
            text = f"<b>{text}</b>"
        if run.italic:
            text = f"<i>{text}</i>"
        if run.underline:
            text = f"<u>{text}</u>"
        parts.append(text)
    joined = "".join(parts).strip()
    if not joined:
        joined = (paragraph.text or "").strip()
    return joined


def _docx_table_to_flowable(table):
    data = []
    for row in table.rows:
        data.append([cell.text.strip() for cell in row.cells])
    if not data:
        return None
    t = Table(data, hAlign="LEFT")
    t.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#c9cede")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eef1fb")),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    return t


def convert_word_to_pdf(input_path, output_path):
    """Convert a .doc/.docx file into a .pdf file.

    Raises ConversionError on invalid or unreadable input.
    """
    if not os.path.isfile(input_path):
        raise ConversionError("Uploaded file could not be found.")

    try:
        document = Document(input_path)
    except PackageNotFoundError as exc:
        raise ConversionError(
            "The uploaded file is not a valid Word (.docx) document."
        ) from exc
    except Exception as exc:  # noqa: BLE001
        raise ConversionError(
            "Could not read the Word document. It may be corrupted."
        ) from exc

    styles = _build_pdf_styles()
    story = []

    doc_pdf = SimpleDocTemplate(
        output_path,
        pagesize=LETTER,
        leftMargin=0.9 * inch,
        rightMargin=0.9 * inch,
        topMargin=0.9 * inch,
        bottomMargin=0.9 * inch,
        title="QuickConvert Document",
    )

    list_buffer = []

    def flush_list_buffer():
        if list_buffer:
            story.append(
                ListFlowable(
                    [ListItem(Paragraph(text, styles["QCBody"])) for text in list_buffer],
                    bulletType="bullet",
                    leftIndent=18,
                )
            )
            story.append(Spacer(1, 8))
            list_buffer.clear()

    body_elements = document.element.body
    table_iter = iter(document.tables)
    processed_tables = set()

    for child in body_elements.iterchildren():
        tag = child.tag.split("}")[-1]

        if tag == "p":
            matching_para = None
            for para in document.paragraphs:
                if para._p is child:
                    matching_para = para
                    break
            if matching_para is None:
                continue

            text = _runs_to_html(matching_para)
            style_name = (matching_para.style.name or "").lower()

            is_list = matching_para.style.name in (
                "List Bullet", "List Bullet 2", "List Number", "List Number 2"
            ) or matching_para._p.pPr is not None and matching_para._p.pPr.numPr is not None

            if not text:
                flush_list_buffer()
                continue

            if is_list:
                list_buffer.append(text)
                continue

            flush_list_buffer()

            if "heading 1" in style_name or style_name == "title":
                story.append(Paragraph(text, styles["QCHeading1"]))
            elif "heading 2" in style_name:
                story.append(Paragraph(text, styles["QCHeading2"]))
            elif "heading 3" in style_name:
                story.append(Paragraph(text, styles["QCHeading3"]))
            else:
                story.append(Paragraph(text, styles["QCBody"]))

        elif tag == "tbl":
            flush_list_buffer()
            try:
                table_obj = next(table_iter)
            except StopIteration:
                table_obj = None
            if table_obj is not None and id(table_obj) not in processed_tables:
                processed_tables.add(id(table_obj))
                flowable = _docx_table_to_flowable(table_obj)
                if flowable is not None:
                    story.append(flowable)
                    story.append(Spacer(1, 12))

    flush_list_buffer()

    if not story:
        story.append(Paragraph("(This document appears to be empty.)", styles["QCBody"]))

    try:
        doc_pdf.build(story)
    except Exception as exc:  # noqa: BLE001
        raise ConversionError("Failed to generate the PDF file.") from exc

    return output_path


# --------------------------------------------------------------------------
# PDF -> Word
# --------------------------------------------------------------------------

def _looks_like_pdf(path):
    try:
        with open(path, "rb") as f:
            header = f.read(5)
        return header.startswith(b"%PDF-")
    except OSError:
        return False


def convert_pdf_to_word(input_path, output_path):
    """Convert a .pdf file into a .docx file.

    Raises ConversionError on invalid or unreadable input.
    """
    if not os.path.isfile(input_path):
        raise ConversionError("Uploaded file could not be found.")

    if not _looks_like_pdf(input_path):
        raise ConversionError("The uploaded file is not a valid PDF document.")

    converter = None
    try:
        converter = PDF2DocxConverter(input_path)
        converter.convert(output_path, start=0, end=None)
    except Exception as exc:  # noqa: BLE001
        raise ConversionError(
            "Could not convert this PDF. It may be scanned, encrypted, or corrupted."
        ) from exc
    finally:
        if converter is not None:
            try:
                converter.close()
            except Exception:  # noqa: BLE001
                pass

    if not os.path.isfile(output_path):
        raise ConversionError("PDF conversion did not produce an output file.")

    return output_path
